#!/usr/bin/env python3
"""
Generate a professional, ATS-safe DOCX resume from a resume JSON structure.

ATS-safe design choices:
- Single column layout
- Standard section headers
- No tables, text boxes, or graphics
- ATS-safe system fonts (Calibri, Georgia, or Consolas — see FONT_MAP)
- Contact info in document body, not header/footer
- Explicit bullet characters (ATS-parseable, no Word list style dependency)
- No em-dashes anywhere (common ATS parse failure)

Usage:
    python generate_resume_docx.py --resume-json '<json>' --output-path 'resume.docx'
    python generate_resume_docx.py --resume-file resume.json --output-path resume.docx
    python generate_resume_docx.py --resume-file resume.json --output-path resume.docx --font serif

Font selection (--font or "font_family" in JSON):
    sans-serif  →  Calibri  (default — clean, modern, highest ATS compatibility)
    serif       →  Georgia  (traditional, slightly warmer — good for consulting/finance)
    mono        →  Consolas (technical feel — unusual for resumes, use sparingly)

JSON schema: see references/output-templates.md in the ats-resume-optimizer skill.
"""

import argparse
import json
import os
import sys

# Maps the CSS-style font family names the user provides to DOCX-safe equivalents.
# Only fonts reliably present in Word/LibreOffice across platforms are listed.
FONT_MAP = {
    "sans-serif": "Calibri",
    "serif":      "Georgia",
    "mono":       "Consolas",
}
DEFAULT_FONT = "Calibri"


def resolve_font(resume_data, cli_font=None):
    """Return the DOCX font name to use, with priority: CLI arg > JSON field > default."""
    key = cli_font or resume_data.get("font_family", "sans-serif")
    return FONT_MAP.get(key, DEFAULT_FONT)


def install_and_import():
    try:
        from docx import Document
        return True
    except ImportError:
        print("Installing python-docx...", file=sys.stderr)
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        return True


def add_right_tab_stop(para, position_inches=6.8):
    """Add a right-aligned tab stop to a paragraph (for title/date layout)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    tabs_elem = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    # position in twentieths of a point (twips)
    tab.set(qn('w:pos'), str(int(position_inches * 72 * 20)))
    tabs_elem.append(tab)
    pPr.append(tabs_elem)


def add_section_heading(doc, text, font=DEFAULT_FONT):
    """Add an ALL CAPS section heading with a bottom border line."""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(3)

    run = para.add_run(text.upper())
    run.font.name = font
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    return para


def add_bullet(doc, text, font=DEFAULT_FONT):
    """
    Add a bullet point using an explicit '•' character with a hanging indent.
    More ATS-reliable than Word's List Bullet style, which some parsers skip entirely.
    """
    from docx.shared import Pt, Inches
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.first_line_indent = Inches(-0.15)
    run = para.add_run(u"\u2022 " + text)
    run.font.name = font
    run.font.size = Pt(10.5)
    return para


def generate_docx(resume_data, output_path, font=None):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    f = font or resolve_font(resume_data)

    doc = Document()

    # Narrow margins for more content space
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    doc.styles['Normal'].font.name = f
    doc.styles['Normal'].font.size = Pt(10.5)

    candidate = resume_data.get("candidate", {})

    # ── NAME (ALL CAPS) ───────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(candidate.get("name", "Your Name").upper())
    name_run.font.name = f
    name_run.font.size = Pt(18)
    name_run.font.bold = True

    # ── TARGET JOB TITLE (subtitle, ALL CAPS) ────────────────────
    # Use candidate["target_title"] if provided, else skip.
    # This should be the role title from the JD, not the current job title.
    target_title = candidate.get("target_title", "")
    if target_title:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(4)
        title_run = title_para.add_run(target_title.upper())
        title_run.font.name = f
        title_run.font.size = Pt(12)
        title_run.font.bold = True

    # ── CONTACT INFO ──────────────────────────────────────────────
    contact_parts = []
    for field in ["email", "phone", "location", "linkedin", "github"]:
        if candidate.get(field):
            contact_parts.append(candidate[field])

    if contact_parts:
        contact_para = doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.paragraph_format.space_after = Pt(6)
        contact_run = contact_para.add_run(" , ".join(contact_parts))
        contact_run.font.name = f
        contact_run.font.size = Pt(10)

    # ── SUMMARY ───────────────────────────────────────────────────
    summary = resume_data.get("summary", "")
    if summary:
        add_section_heading(doc, "Summary", font=f)
        summary_para = doc.add_paragraph()
        summary_para.paragraph_format.space_after = Pt(4)
        summary_run = summary_para.add_run(summary)
        summary_run.font.name = f
        summary_run.font.size = Pt(10.5)

    # ── EXPERIENCE ────────────────────────────────────────────────
    # Each job entry is kept together on the same page using keep_with_next.
    # All paragraphs in an entry chain to the next via keep_with_next = True,
    # except the last bullet, which breaks the chain so entries don't fuse.
    experience = resume_data.get("experience", [])
    if experience:
        add_section_heading(doc, "Experience", font=f)
        for job in experience:
            bullets = job.get("bullets", [])

            # Build date string
            date_str = job.get("start_date", "")
            if job.get("end_date"):
                date_str += f" - {job['end_date']}"

            # Line 1: Job Title (bold, left) + Date (right-aligned via tab stop)
            job_header = doc.add_paragraph()
            job_header.paragraph_format.space_before = Pt(8)
            job_header.paragraph_format.space_after = Pt(0)
            job_header.paragraph_format.keep_with_next = True  # chain to company line
            add_right_tab_stop(job_header)

            title_run = job_header.add_run(job.get("title", ""))
            title_run.font.name = f
            title_run.font.size = Pt(11)
            title_run.font.bold = True

            if date_str:
                tab_run = job_header.add_run("\t")
                tab_run.font.name = f
                date_run = job_header.add_run(date_str)
                date_run.font.name = f
                date_run.font.size = Pt(10.5)

            # Line 2: Company name (italic) + optional location
            company_parts = []
            if job.get("company"):
                company_parts.append(job["company"])
            if job.get("location"):
                company_parts.append(job["location"])

            if company_parts:
                company_para = doc.add_paragraph()
                company_para.paragraph_format.space_before = Pt(0)
                company_para.paragraph_format.space_after = Pt(3)
                company_para.paragraph_format.keep_with_next = True  # chain to first bullet
                company_run = company_para.add_run(", ".join(company_parts))
                company_run.font.name = f
                company_run.font.size = Pt(10.5)
                company_run.font.italic = True

            for i, bullet in enumerate(bullets):
                para = add_bullet(doc, bullet, font=f)
                # Chain all bullets to the next EXCEPT the last one.
                # The last bullet ends the keep_with_next chain so this entry
                # doesn't bleed into the next entry's page-break logic.
                if i < len(bullets) - 1:
                    para.paragraph_format.keep_with_next = True

    # ── SKILLS ────────────────────────────────────────────────────
    # Categorized skills are better for ATS keyword parsing than a flat list.
    # The entire skills section is chained with keep_with_next so it never
    # splits mid-section across a page break.
    skills = resume_data.get("skills", [])
    languages = resume_data.get("languages", [])
    all_skill_rows = list(skills) + (["__languages__"] if languages else [])

    if all_skill_rows:
        add_section_heading(doc, "Skills", font=f)
        for i, skill_group in enumerate(all_skill_rows):
            is_last = (i == len(all_skill_rows) - 1)

            skills_para = doc.add_paragraph()
            skills_para.paragraph_format.space_before = Pt(1)
            skills_para.paragraph_format.space_after = Pt(1)
            if not is_last:
                skills_para.paragraph_format.keep_with_next = True

            if skill_group == "__languages__":
                label = skills_para.add_run("Languages: ")
                label.font.name = f
                label.font.size = Pt(10.5)
                label.font.bold = True
                val = skills_para.add_run(", ".join(languages))
                val.font.name = f
                val.font.size = Pt(10.5)
            else:
                category_run = skills_para.add_run(f"{skill_group.get('category', 'Skills')}: ")
                category_run.font.name = f
                category_run.font.size = Pt(10.5)
                category_run.font.bold = True
                items_run = skills_para.add_run(", ".join(skill_group.get("items", [])))
                items_run.font.name = f
                items_run.font.size = Pt(10.5)

    # ── EDUCATION ─────────────────────────────────────────────────
    education = resume_data.get("education", [])
    if education:
        add_section_heading(doc, "Education", font=f)
        for edu in education:
            # Line 1: Degree (bold) + graduation year (right-aligned)
            edu_header = doc.add_paragraph()
            edu_header.paragraph_format.space_before = Pt(4)
            edu_header.paragraph_format.space_after = Pt(0)
            add_right_tab_stop(edu_header)

            degree_run = edu_header.add_run(edu.get("degree", ""))
            degree_run.font.name = f
            degree_run.font.size = Pt(11)
            degree_run.font.bold = True

            if edu.get("graduation"):
                edu_header.add_run("\t")
                grad_run = edu_header.add_run(edu["graduation"])
                grad_run.font.name = f
                grad_run.font.size = Pt(10.5)

            # Line 2: Institution (italic)
            if edu.get("institution"):
                inst_para = doc.add_paragraph()
                inst_para.paragraph_format.space_before = Pt(0)
                inst_para.paragraph_format.space_after = Pt(2)
                details = [edu["institution"]]
                if edu.get("gpa"):
                    details.append(f"GPA: {edu['gpa']}")
                if edu.get("notes"):
                    details.append(edu["notes"])
                inst_run = inst_para.add_run(", ".join(details))
                inst_run.font.name = f
                inst_run.font.size = Pt(10.5)
                inst_run.font.italic = True

    # ── CERTIFICATIONS ────────────────────────────────────────────
    certifications = resume_data.get("certifications", [])
    if certifications:
        add_section_heading(doc, "Certifications", font=f)
        for cert in certifications:
            cert_para = doc.add_paragraph()
            cert_para.paragraph_format.space_before = Pt(2)
            cert_para.paragraph_format.space_after = Pt(1)
            add_right_tab_stop(cert_para)

            cert_run = cert_para.add_run(cert.get("name", ""))
            cert_run.font.name = f
            cert_run.font.size = Pt(10.5)
            cert_run.font.bold = True

            parts = []
            if cert.get("issuer"):
                parts.append(cert["issuer"])
            if cert.get("date"):
                cert_para.add_run("\t")
                date_run = cert_para.add_run(cert["date"])
                date_run.font.name = f
                date_run.font.size = Pt(10.5)
            elif parts:
                detail_run = cert_para.add_run(f", {', '.join(parts)}")
                detail_run.font.name = f
                detail_run.font.size = Pt(10.5)

    # ── PROJECTS ─────────────────────────────────────────────────
    projects = resume_data.get("projects", [])
    if projects:
        add_section_heading(doc, "Projects", font=f)
        for project in projects:
            proj_para = doc.add_paragraph()
            proj_para.paragraph_format.space_before = Pt(4)
            proj_para.paragraph_format.space_after = Pt(1)

            proj_run = proj_para.add_run(project.get("name", ""))
            proj_run.font.name = f
            proj_run.font.size = Pt(11)
            proj_run.font.bold = True

            if project.get("tech"):
                tech_run = proj_para.add_run(f" ({project['tech']})")
                tech_run.font.name = f
                tech_run.font.size = Pt(10.5)

            for bullet in project.get("bullets", []):
                add_bullet(doc, bullet, font=f)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f"Resume saved: {output_path}")


def main():
    parser = argparse.ArgumentParser(description="Generate ATS-safe DOCX resume")
    parser.add_argument("--resume-json", help="Resume data as JSON string")
    parser.add_argument("--resume-file", help="Path to resume JSON file")
    parser.add_argument("--output-path", required=True, help="Output .docx file path")
    parser.add_argument(
        "--font",
        choices=list(FONT_MAP.keys()),
        default=None,
        help="Font family override: sans-serif (Calibri, default), serif (Georgia), mono (Consolas). "
             "Also readable from resume JSON 'font_family' field.",
    )
    args = parser.parse_args()

    if not args.resume_json and not args.resume_file:
        print("Error: provide --resume-json or --resume-file", file=sys.stderr)
        sys.exit(1)

    if args.resume_file:
        with open(args.resume_file) as fh:
            resume_data = json.load(fh)
    else:
        resume_data = json.loads(args.resume_json)

    install_and_import()
    font = resolve_font(resume_data, cli_font=args.font)
    generate_docx(resume_data, args.output_path, font=font)


if __name__ == "__main__":
    main()
